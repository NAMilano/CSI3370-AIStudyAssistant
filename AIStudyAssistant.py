import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from datetime import datetime
import requests
import os
import itertools
import random
import json
    
# select and load a locally saved document 
class LoadDocument:
    def selectDocument(self):
        # get file path - SUPPORTED FILE TYPES: docx, txt, pdf
        path = filedialog.askopenfilename(filetypes=[("Word Document", "*.docx"), ("Text File", "*.txt"), ("PDF Document", "*.pdf")])

        # check to make sure a file was selected - if not exit
        if not path:
            return None, None
        
        # get the name of the file
        fileName = os.path.basename(path)
        return fileName, path


# parse a word doc, txt file, or PDF file compiling all the contents into a single string
class ParseDocument:
    def parse(self, path):
        if not path:
            return None
        try:
            # get the file type of the selected document
            fileType = path.split(".")[-1].lower()

            # word document parsing
            if fileType == "docx":
                wordDoc = Document(path)
                # combine all the contents of the document into one string
                contents = '\n'.join(sections.text for sections in wordDoc.paragraphs if sections.text)
                
            
            # # text file parsing
            # elif fileType == "txt":
            #     # parsing code here
            #    
            
            # # pdf file parsing
            # elif fileType == "pdf":
            #     #parsing code here
            #   
    
            return contents

        # error message exception catch
        except Exception as e:
            messagebox.showerror("Error", f"Error loading document: {e}")
            return None

    

# handles the Gemini API call 
class GeminiServices:
    # calls the Gemini API using a HTTP request and then receives the response in a json object
    def call(self, prompt):
        headers = {"Content-Type": "application/json"}
        data = {"contents": [{"parts": [{"text": prompt}]}], 
                 # limit output to 512 tokens ~680 words
                 # temperature set to 0.7 to balance randomness ensuring accuracy levels but allows for some creativity
                "generationConfig": {"maxOutputTokens": 512, "temperature": 0.7}
        }
        try:
            response = requests.post("https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key=AIzaSyC1hocc2oAhNvwTyPAmITPIMMbgocHWihc", headers = headers, json = data)
            response.raise_for_status()
            # get the generated AI answer from the returned json object
            geminiResponse = response.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
            return geminiResponse
        except Exception as e:
            raise Exception(f"Error calling Gemini API: {e}")
        



# creates quiz questions for the selected document
# a max of 5 quiz questions but depending on the contents of the document fewer can be presented
# a keyword from the question is replaced with _______ for a fill-in-the-blank quiz experience
# Example: 
#   Question: The ________ is the main heat source for the Earth.
#   Answer: Sun
class QuizQuestions:
    # get the GeminiServices instance from the GUI class
    def __init__(self, geminiCall):
        self.geminiCall = geminiCall

    #
    def generate(self, contents):
        # check that contents string is populated
        if not contents:
            raise ValueError("Error: Selected document has no contents")
        
        # prompt used to guide the AI to create the quiz questions - Specifies how many, what we want, and how to format the answer
        prompt = f"Generate 5 fill-in-the-blank quiz questions from this text. Someone viewing these questions must be able to know the answer solely based on the information in the text they are being generated from. Format each as '(# of question). [question]\tAnswer: [answer]'\n. Example response: 1. The _____ is the main heat source for the Earth.\tAnswer: Sun\n. Here is the text: {contents}"
        raw = self.geminiCall.call(prompt)

        questions = []

        for ln in raw.splitlines():
            if "\t" in ln:
                q, a = ln.split("\t", 1)
                questions.append((q.strip(), a.strip()))

        if not questions:
            raise RuntimeError("Gemini did not return valid quiz questions. Please try again.")
        
        return questions
    
class Flashcards:
    def __init__(self, geminiCall):
        self.geminiCall = geminiCall  # keep the same attribute name

    def generate(self, contents, n_cards=15):
        if not contents:
            raise ValueError("Error: Selected document has no contents")

        # TSV prompt
        prompt = (
            f"Create exactly {n_cards} flash-cards from the text below. Return ONLY plain text with one card per line, using a TAB between the question and the answer. Example: Are bananas berries?\tYes\n----- TEXT START -----\n"
            f"{contents}\n----- TEXT END -----"
        )
        raw = self.geminiCall.call(prompt)

        cards = []

        for ln in raw.splitlines():
            if "\t" in ln:
                q, a = ln.split("\t", 1)
                cards.append((q.strip(), a.strip()))

        if not cards:
            raise RuntimeError("Gemini did not return TSV flashcards..")
        return cards
    
class TopicExtractor:
    def __init__(self, geminiCall):
        self.geminiCall = geminiCall

    def analyze(self, contents):
        if not contents:
            raise ValueError("Error: Selected document has no contents")
        
        # prompt to extract key topics with clear formatting
        prompt = f"""Analyze the following document and extract the most important key topics and concepts. 
        Present them as a clear, organized list. Focus on the main themes, important terms, and core concepts that a student should understand.
        Format your response EXACTLY like this example:

        • Main Topic 1 - Brief explanation
        • Main Topic 2 - Brief explanation  
        • Key Concept 1 - What it means
        • Important Term - Definition or context
        • Core Principle - Why it matters

        Use bullet points (•) and keep each point concise but informative. Limit to 8-12 key topics maximum. Do not include any additional text or explanations outside the requested format.
        
        Document text: {contents}"""
        
        topics = self.geminiCall.call(prompt)
        
        if not topics:
            raise RuntimeError("Gemini did not return key topics. Please try again.")
        
        return topics


class FlashcardViewer(tk.Toplevel):
    #class for flashcard viewer 

    def __init__(self, parent, cards):
        super().__init__(parent)
        self.title("Flash‑Card Review")
        self.geometry("500x260")

        # randomize the order of cards then cycles the same.
        random.shuffle(cards)
        self.deck = itertools.cycle(cards)

        # tkinter variables for question and answer
        self.q_var = tk.StringVar()
        self.a_var = tk.StringVar()

        # layout for the popup questions
        tk.Label(self, textvariable=self.q_var, wraplength=460,
                 font=("Helvetica", 15, "bold")).pack(pady=(25, 10))
        tk.Label(self, textvariable=self.a_var, wraplength=460,
                 font=("Helvetica", 14)).pack(pady=(0, 15))

        bar = tk.Frame(self)
        bar.pack()
        tk.Button(bar, text="Show Answer",
                  command=self.show_answer).pack(side=tk.LEFT, padx=6)
        tk.Button(bar, text="Next Card",
                  command=self.next_card).pack(side=tk.LEFT, padx=6)

        self.current = None
        self.next_card()          # shows the first q

    # helper functions 
    def next_card(self):
       # Load the next card and hide its answer
        self.current = next(self.deck)
        self.q_var.set(self.current[0])
        self.a_var.set("")

    def show_answer(self):
        # shows answer
        self.a_var.set(self.current[1])
    

class PromptController:
    def __init__(self):
        self.geminiCall = GeminiServices()
        self.quiz = QuizQuestions(self.geminiCall)
        self.flash = Flashcards(self.geminiCall)
        self.extractor = TopicExtractor(self.geminiCall)


    def generateQuizQuestions(self, contents):
        questions = self.quiz.generate(contents)
        return questions
    
    def generateFlashcards(self, contents, n_cards=15):
        return self.flash.generate(contents, n_cards)
    
    def extractKeyTopics(self, contents):
        return self.extractor.analyze(contents)
    
# Pomodoro timer class
# Allows user to set work and break durations and track Pomodoro sessions throughout the day
# Includes progress bar, session counter, goal tracker, and persistent session storage
class PomodoroTimer(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("Pomodoro Timer")
        self.geometry("320x400")

        # initialize session tracking file and today's date
        self.data_file = "pomodoro_data.json"
        self.today = datetime.now().strftime("%Y-%m-%d")
        self.session_count = self.load_session_count()

        # initialize control flags and variables
        self.is_running = False
        self.timer_id = None

        # timer state variables
        self.work_duration = tk.IntVar(value=25)      # default 25 min work
        self.break_duration = tk.IntVar(value=5)      # default 5 min break
        self.goal = tk.IntVar(value=4)                # default 4 Pomodoros goal
        self.time_left = tk.StringVar(value="00:00")
        self.current_phase = "Work"

        # input fields: user sets durations and daily goal
        tk.Label(self, text="Work Duration (minutes):").pack(pady=3)
        tk.Entry(self, textvariable=self.work_duration).pack()

        tk.Label(self, text="Break Duration (minutes):").pack(pady=3)
        tk.Entry(self, textvariable=self.break_duration).pack()

        tk.Label(self, text="Daily Goal (Pomodoros):").pack(pady=3)
        tk.Entry(self, textvariable=self.goal).pack()

        # label to show current phase (Work or Break)
        self.phase_label = tk.Label(self, text=f"Phase: {self.current_phase}", font=("Helvetica", 12, "bold"))
        self.phase_label.pack(pady=10)

        # large display for countdown
        self.timer_label = tk.Label(self, textvariable=self.time_left, font=("Helvetica", 28))
        self.timer_label.pack(pady=5)

        # progress bar to visually show timer progress
        self.progress = ttk.Progressbar(self, length=250, mode='determinate', maximum=100)
        self.progress.pack(pady=5)

        # start, stop, and reset buttons
        button_frame = tk.Frame(self)
        button_frame.pack(pady=10)
        tk.Button(button_frame, text="Start", command=self.start_timer).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="Stop", command=self.stop_timer).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="Reset", command=self.reset_timer).pack(side=tk.LEFT, padx=5)

        # session status display (completed & goal)
        self.session_label = tk.Label(self, text=f"Pomodoros completed: {self.session_count}")
        self.session_label.pack()
        self.goal_label = tk.Label(self, text=self.goal_progress_text())
        self.goal_label.pack()

    # load the session count from the saved JSON file if the date matches today
    def load_session_count(self):
        if os.path.exists(self.data_file):
            with open(self.data_file, "r") as f:
                data = json.load(f)
                if data.get("date") == self.today:
                    return data.get("count", 0)
        return 0

    # save session count and date to a JSON file for persistence
    def save_session_count(self):
        with open(self.data_file, "w") as f:
            json.dump({"date": self.today, "count": self.session_count}, f)

    # generate the text showing goal progress
    def goal_progress_text(self):
        return f"Goal progress: {self.session_count}/{self.goal.get()} {'✅' if self.session_count >= self.goal.get() else ''}"

    # starts the Pomodoro timer
    def start_timer(self):
        if self.is_running:
            return
        self.is_running = True
        duration = self.work_duration.get() * 60
        self.run_phase(duration, "Work")

    # stops the timer (does not reset)
    def stop_timer(self):
        if self.timer_id:
            self.after_cancel(self.timer_id)
        self.is_running = False

    # resets timer display and progress bar to current phase's full duration
    def reset_timer(self):
        self.stop_timer()
        duration = self.work_duration.get() * 60 if self.current_phase == "Work" else self.break_duration.get() * 60
        mins, secs = divmod(duration, 60)
        self.time_left.set(f"{mins:02}:{secs:02}")
        self.progress['value'] = 0
        self.phase_label.config(text=f"Phase: {self.current_phase}")

    # begins countdown for given phase (Work or Break)
    def run_phase(self, duration, phase):
        self.current_phase = phase
        self.total_phase_time = duration
        self.phase_label.config(text=f"Phase: {self.current_phase}")
        self.countdown(duration)

    # updates countdown every second, adjusts progress bar
    def countdown(self, remaining):
        mins, secs = divmod(remaining, 60)
        self.time_left.set(f"{mins:02}:{secs:02}")

        percent = (1 - remaining / self.total_phase_time) * 100
        self.progress['value'] = percent

        if remaining > 0:
            self.timer_id = self.after(1000, self.countdown, remaining - 1)
        else:
            self.is_running = False
            self.notify_phase_end()

            if self.current_phase == "Work":
                # increment session count after a work session
                self.session_count += 1
                self.session_label.config(text=f"Pomodoros completed: {self.session_count}")
                self.goal_label.config(text=self.goal_progress_text())
                self.save_session_count()

                # show popup if user reached their daily goal
                if self.session_count == self.goal.get():
                    messagebox.showinfo("Goal Reached!", "🎉 You've completed your Pomodoro goal for the day!")

                # begin break phase
                self.run_phase(self.break_duration.get() * 60, "Break")

            elif self.current_phase == "Break":
                # begin work phase
                self.run_phase(self.work_duration.get() * 60, "Work")

    # shows popup and beeps at the end of each phase
    def notify_phase_end(self):
        try:
            self.bell()
        except:
            pass
        if self.current_phase == "Work":
            messagebox.showinfo("Break Time!", "Time for a break!")
        else:
            messagebox.showinfo("Work Time!", "Back to work!")


# GUI class for the AI study assistant
class StudyAssistantGUI:
    def __init__(self, root):
        # setup GUI main window
        self.root = root
        self.root.title("AI Study Assistant - Powered by Gemini")
        self.root.geometry("700x500")
        self.contents = ""

        # initialize class objects
        self.loadDoc = LoadDocument()
        self.parseDoc = ParseDocument()
        self.promptCon = PromptController()
        self.flashcardWindow = None
        self.pomodoroWindow = None
        self.viewAnswersButton = None

        # select document button
        tk.Button(self.root, text="Select Document", command = self.selectDocument).pack(pady=5)

        # status label to show file name of the loaded document
        self.status_label = tk.Label(self.root, text="No document selected")
        self.status_label.pack(pady=5) 

        # function buttons - summarization, keyTopics, flashcards, QuizQuestions, pomodoroTimer
        self.buttons = tk.Frame(self.root)
        self.buttons.pack(pady=5)
        tk.Button(self.buttons, text="Generate Quiz", command=self.runQuiz).pack(side=tk.LEFT, pady=5)
        tk.Button(self.buttons, text="Generate Flashcards", command=self.runFlashcards).pack(side=tk.LEFT, padx=5)
        tk.Button(self.buttons, text="Get Key Topics", command=self.getKeyTopics).pack(side=tk.LEFT, padx=5) 
        tk.Button(self.buttons, text="Pomodoro Timer", command=self.runPomodoro).pack(side=tk.LEFT, padx=5)

        # text output area 
        self.outputArea = tk.Text(self.root, height=20, width=80, wrap="word")
        self.outputArea.pack(side=tk.BOTTOM, pady=15)



    # select and load locally saved document
    # updates the UI: shows confirmation message in outputArea - shows filename under "Select Document" button
    def selectDocument(self):
        fileName, path = self.loadDoc.selectDocument()
        if fileName and path:
            self.contents = self.parseDoc.parse(path)
            # clear outputArea
            self.outputArea.delete(1.0, tk.END)
            # display success message
            self.outputArea.insert(tk.END, f"Document '{fileName}' loaded.\n")
            # upadate status label to show filename under "Select Document" button
            self.status_label.config(text=fileName)
    

    # generate quiz questions
    # updates the UI with the generated quiz questions
    def runQuiz(self):
        # check for a loaded document 
        if not self.contents:
            # error message when no document is loaded
            messagebox.showwarning("No document available", "Please select a document.")
            return
        
        # update output area with a processing message
        self.outputArea.delete(1.0, tk.END)
        self.outputArea.insert(tk.END, "Generating Quiz Questions...")
        self.root.update()

        try:
            # generate questions
            self.questions = self.promptCon.generateQuizQuestions(self.contents)
            # clear output area
            self.outputArea.delete(1.0, tk.END)

            for q, a in self.questions:
                self.outputArea.insert(tk.END, q + "\n\n")
            
            tempButton = tk.Frame(self.root)
            tempButton.pack(pady=5) 

            self.viewAnswersButton = tk.Button(tempButton, text="View Quiz Answers", command=self.showAnswers)
            self.viewAnswersButton.pack(side=tk.BOTTOM, padx=5)


        except Exception as e:
            self.outputArea.delete(1.0, tk.END)
            # display error message
            self.outputArea.insert(tk.END, f"Error: {e}\n")


    def showAnswers(self):
        self.outputArea.insert(tk.END, "\n")
        for q, a in self.questions:
            self.outputArea.insert(tk.END, a + "\n")

        # destroy the view answers button after displaying the results
        self.viewAnswersButton.destroy()
        self.viewAnswersButton = None


    def runFlashcards(self):
        if not self.contents:
            messagebox.showwarning("No document available", "Please select a document.")
            return

        # show processing message
        self.outputArea.delete("1.0", tk.END)
        self.outputArea.insert(tk.END, "Generating flash-cards...")
        self.root.update()

        try:
            #flashcard method
            cards = self.promptCon.generateFlashcards(self.contents)  # returns list of (Q, A)
            self.outputArea.delete("1.0", tk.END)
            self.outputArea.insert(tk.END, f"{len(cards)} cards generated – launching viewer…")

            # check if the window is already created/open - if it is close it
            if self.flashcardWindow is not None and self.flashcardWindow.winfo_exists():
                self.flashcardWindow.destroy()

            # create flashcard UI element
            self.flashcardWindow = FlashcardViewer(self.root, cards) 
        except Exception as e:
            self.outputArea.delete("1.0", tk.END)
            self.outputArea.insert(tk.END, f"Error: {e}")

    def getKeyTopics(self):
        # check for a loaded document 
        if not self.contents:
            # error message when no document is loaded
            messagebox.showwarning("No document available", "Please select a document.")
            return
        
        # update output area with a processing message
        self.outputArea.delete(1.0, tk.END)
        self.outputArea.insert(tk.END, "Extracting key topics from document...")
        self.root.update()

        try:
            # extract key topics
            topics = self.promptCon.extractKeyTopics(self.contents)
            # clear output area and display topics with nice formatting
            self.outputArea.delete(1.0, tk.END)
            
            # Add attractive header
            self.outputArea.insert(tk.END, "📚 KEY TOPICS & CONCEPTS\n")
            self.outputArea.insert(tk.END, "=" * 60 + "\n\n")
            
            # Process and format the topics nicely
            formatted_topics = self.formatTopics(topics)
            self.outputArea.insert(tk.END, formatted_topics)
            
            # Add footer
            self.outputArea.insert(tk.END, "\n\n" + "=" * 60)
            self.outputArea.insert(tk.END, "\n💡 Tip: Use these topics to guide your study focus!")
            
        except Exception as e:
            self.outputArea.delete(1.0, tk.END)
            # display error message
            self.outputArea.insert(tk.END, f"Error: {e}\n")       

    # Helper method to format topics nicely
    def formatTopics(self, topics):
        lines = topics.split('\n')
        formatted = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                formatted += "\n"
                continue
                
            # If line doesn't start with bullet, add one
            if line and not line.startswith('•') and not line.startswith('-') and not line.startswith('*'):
                if line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                    # Convert numbered list to bullets
                    line = '• ' + line[2:].strip()
                else:
                    line = '• ' + line
            elif line.startswith('-') or line.startswith('*'):
                # Convert other bullet types to consistent format
                line = '• ' + line[1:].strip()
            
            formatted += line + "\n"
        
        return formatted

    def runPomodoro(self):
        # check if the window is already created/open - if it is close it
        if self.pomodoroWindow is not None and self.pomodoroWindow.winfo_exists():
            self.pomodoroWindow.lift()
            return

        # create the Pomodoro timer window
        self.pomodoroWindow = PomodoroTimer(self.root)

if __name__ == "__main__":
    root = tk.Tk()
    app = StudyAssistantGUI(root)
    root.mainloop()
            
        









